from __future__ import annotations

import json
import tempfile
from datetime import date
from pathlib import Path

from app.rag.document_loader import (
    DOCXDocumentLoader,
    DocumentMetadataError,
    OCRRequiredError,
    docx_metadata_sidecar_path,
)
from app.rag.policy_chunker import chunk_policy_directory
from app.rag.policy_context import build_policy_context
from app.rag.policy_parser import parse_policy_directory
from app.rag.policy_retriever import PolicyRetrievalResult, PolicyRetriever
from app.schemas.policy import SecurityLevel
from app.security import PolicyAccessContext, TrustedIdentitySource


class _OfflineEmbeddingProvider:
    @property
    def dimension(self) -> int:
        return 1

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [[1.0] for _ in texts]

    def embed_query(self, text: str) -> list[float]:
        del text
        return [1.0]


def _metadata_text() -> str:
    return """document_id: DOCX_VERIFY_001
document_type: policy
title: DOCX 离线验收制度
version: "1.0"
status: effective
issuing_department: 技术部
effective_date: 2026-01-01
allowed_departments:
  - ALL
allowed_roles:
  - EMPLOYEE
security_level: internal
region: 中国大陆
"""


def _write_sidecar(path: Path) -> None:
    docx_metadata_sidecar_path(path).write_text(_metadata_text(), encoding="utf-8")


def _write_docx(path: Path, *, with_text: bool = True) -> None:
    from docx import Document

    document = Document()
    if with_text:
        document.add_paragraph("DOCX 离线验收制度", style="Title")
        document.add_paragraph("第一章 总则", style="Heading 1")
        document.add_paragraph("第一条 适用范围", style="Heading 2")
        document.add_paragraph("本制度适用于全体员工。")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "材料"
        table.cell(0, 1).text = "要求"
        table.cell(1, 0).text = "申请表"
        table.cell(1, 1).text = "必须提交"
        document.add_paragraph("第二条 保密要求", style="Heading 2")
        document.add_paragraph("员工不得泄露企业内部信息。")
    document.save(path)


def run_verification() -> dict[str, object]:
    """Verify DOCX extraction, provenance, sidecar trust, and security offline."""

    with tempfile.TemporaryDirectory(prefix="enterprise-policy-docx-") as temporary_directory:
        root = Path(temporary_directory)
        source = root / "policy.docx"
        _write_docx(source)
        _write_sidecar(source)

        documents = parse_policy_directory(root)
        chunks = chunk_policy_directory(root)
        document = documents[0]
        context = build_policy_context([PolicyRetrievalResult(chunk=chunks[1], score=1.0)])

        access_context = PolicyAccessContext(
            employee_id="DOCX-VERIFY-001",
            department="技术部",
            roles=("EMPLOYEE",),
            security_clearance=SecurityLevel.INTERNAL,
            region="中国大陆",
            identity_source=TrustedIdentitySource.TEST_FIXTURE,
        )
        retriever = PolicyRetriever(
            embedding_provider=_OfflineEmbeddingProvider(),
            chunks=chunks,
        ).restrict(access_context, as_of_date=date(2026, 8, 20))
        secured_results = retriever.search("保密要求", top_k=2)

        image_only_source = root / "image-only.docx"
        _write_docx(image_only_source, with_text=False)
        _write_sidecar(image_only_source)
        ocr_required = False
        try:
            DOCXDocumentLoader().load(image_only_source)
        except OCRRequiredError:
            ocr_required = True

        missing_sidecar_source = root / "missing-sidecar.docx"
        _write_docx(missing_sidecar_source)
        missing_sidecar_rejected = False
        try:
            DOCXDocumentLoader().load(missing_sidecar_source)
        except DocumentMetadataError:
            missing_sidecar_rejected = True

        citation = context.citations[0]
        checks = {
            "docx_discovered_without_sidecar_as_document": len(documents) == 1,
            "trusted_sidecar_metadata_applied": (
                document.document_id == "DOCX_VERIFY_001"
                and document.metadata_source_path == docx_metadata_sidecar_path(source)
            ),
            "paragraphs_and_styles_extracted": (
                document.source_loader_name == "python-docx"
                and "## 第一章 总则" in document.content
            ),
            "tables_preserved_in_document_order": ("| 申请表 | 必须提交 |" in chunks[0].content),
            "existing_parser_and_chunker_reused": (
                len(chunks) == 2
                and [chunk.article_label for chunk in chunks] == ["第一条", "第二条"]
            ),
            "block_provenance_preserved": (
                chunks[0].source_block_start == 3
                and chunks[0].source_block_end == 5
                and citation.source_block_start == 6
                and citation.source_block_end == 7
            ),
            "authorization_still_precedes_similarity": (
                retriever.allowed_chunk_count == 2 and len(secured_results) == 2
            ),
            "image_only_docx_requests_future_ocr": ocr_required,
            "missing_sidecar_is_rejected": missing_sidecar_rejected,
        }

        return {
            "schema_version": "1.0",
            "phase": 24,
            "passed": all(checks.values()),
            "parser": "python-docx-native-content",
            "document_count": len(documents),
            "source_block_count": document.source_block_count,
            "chunk_count": len(chunks),
            "ocr_executed": False,
            "network_calls": False,
            "model_calls": False,
            "checks": checks,
        }


def main() -> int:
    try:
        report = run_verification()
    except (ImportError, OSError, ValueError) as exc:
        report = {
            "schema_version": "1.0",
            "phase": 24,
            "passed": False,
            "error_type": type(exc).__name__,
            "message": str(exc),
        }
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
